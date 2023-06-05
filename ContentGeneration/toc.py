from docx.shared import RGBColor
from docx.oxml import OxmlElement, ns
from docx.shared import Inches, Pt, Cm
from docx.oxml.ns import qn


def toc(vAR_new_doc):
    paragraph1 = vAR_new_doc.add_paragraph()
    toc = paragraph1.add_run("\t\t\t\t Table of Contents \t")
    toc.bold = True
    toc.font.size = Pt(15)
    toc.font.color.rgb = RGBColor(0, 0, 153)
    run = paragraph1.add_run()
    fldChar = OxmlElement('w:fldChar')  # creates a new element
    fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
    instrText.text = 'TOC \\o "1-2" \\h \\z \\u'   # change 1-3 depending on heading levels you need

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:updateFields')
    fldChar3.set(qn('w:val'), 'true')
    ##fldChar3.text = "Right-click to update field."
    fldChar2.append(fldChar3)

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)
    p_element = paragraph1._p

    vAR_new_doc.add_page_break()