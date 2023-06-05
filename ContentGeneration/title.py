from docx.shared import RGBColor
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
  
def title(vAR_new_doc,vAR_prompt_file):  
    vAR_new_doc.add_paragraph("")
    vAR_new_doc.add_paragraph("")
    vAR_new_doc.add_paragraph("")
    vAR_new_doc.add_paragraph("")
    vAR_new_doc.add_paragraph("")
    vAR_new_doc.add_paragraph("")
    vAR_new_doc.add_paragraph("")
    vAR_new_doc.add_paragraph("")
    vAR_new_doc.add_paragraph("")
    p=vAR_new_doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(vAR_prompt_file)
    r.font.size = Pt(24)
    r.bold = True
    r.underline = True
    r.font.color.rgb = RGBColor(0, 0, 153)
    vAR_new_doc.add_page_break()