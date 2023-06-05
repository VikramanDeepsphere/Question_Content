import streamlit as st
from streamlit_option_menu import option_menu
import os
import io
import source.title_1 as head
import openai
import source.GPT_3 as gpt
import docx
from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, ns
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import ContentGeneration.title as title
import ContentGeneration.toc as toc
import ContentGeneration.header as header
import ContentGeneration.pagenum as page_num
import GPTTurbo.contentgen35 as GPTTurbo

def content_gen():
    openai.api_key = os.environ["API_KEY"]
    head.title()
    st.markdown("<p style='text-align: center; color: black; font-size:20px;'><span style='font-weight: bold'>Problem Statement: </span>ChatGPT Powered Business Applications for Learning Material Generation  </p>", unsafe_allow_html=True)
    st.markdown("<hr style=height:2.5px;background-color:gray>",unsafe_allow_html=True)
    w1,col1,col2,w2=st.columns((0.3,2.5,2.5,0.5))
    with col1:
        st.markdown("### ")
        st.write('# ')
        st.markdown("<p style='text-align: left; color: black; font-size:20px;'><span style='font-weight: 600'>Select Model</span></p>", unsafe_allow_html=True)
        st.write("### ")
        st.markdown("")
    with col2:
        st.markdown("## ")
        vAR_Model = st.radio("",["GPT-3","GPT-3.5"],horizontal=True)
        if vAR_Model != "Select":
            with col1:
                st.markdown("<p style='text-align: left; color: black; font-size:20px;'><span style='font-weight: 600'>Enter Model Input (Content Topic)</span></p>", unsafe_allow_html=True)
            vAR_prompt_file_1 = st.text_input("",placeholder="Hit ENTER after giving topic",key="Clear1")
            vAR_prompt_file_1=vAR_prompt_file_1.strip().split(" ")
            vAR_prompt=[]
            for i in vAR_prompt_file_1:
                if i!="":
                    vAR_prompt.append(i)
            vAR_prompt_file = " ".join(vAR_prompt)
            #Creating a new word document.
            vAR_new_doc = docx.Document()
            title.title(vAR_new_doc,vAR_prompt_file)
            toc.toc(vAR_new_doc)
            header.header(vAR_new_doc)
            page_num.put_page_num(vAR_new_doc)
            if vAR_Model == "GPT-3":
                #Giving a not static input and asking gpt3 to generate 25 random topics from the input.  
                vAR_topics_list = ["Give 25 subtopics for" + "'"+vAR_prompt_file+"'"+" from basics"]

                #Defining a prompt to instruct the gpt3 to generate content for each topics about 200 words. 
                vAR_elaborate_prompt = "Elaborate the topic for 150 to 200 words without conclusion for"

                vAR_conclusion = "Write a conclusion for " + vAR_prompt_file
                with col2:
                    if vAR_prompt_file !="":
                        st.markdown("")
                        button_placeholder = st.empty()
                        button_clicked = False
                        key=0
                        while not button_clicked:
                            key=key+1
                            button_clicked = button_placeholder.button('Generate Content',key=key)
                            break
                        if button_clicked:
                            if vAR_prompt_file !="":
                                button_placeholder.empty()
                                for one_sentence in vAR_topics_list:
                                    with col2:
                                        st.info("Generating Subtopics if the Given Topic is too Broad")
                                    vAR_GPT_response_nonstatic = gpt.generate_response(one_sentence)
                                    with col2:
                                        st.info("Getting Model Response for Main and Sub Topics")
                                    vAR_splitted_response = vAR_GPT_response_nonstatic.split("\n")
                                    with col2:
                                        st.info("Content Summarization")
                                    for j in vAR_splitted_response:
                                        if j=="":
                                            vAR_splitted_response.remove(j)
                                    for topics in vAR_splitted_response:
                                        if topics != "":
                                            vAR_joined_topics=vAR_elaborate_prompt+topics
                                            response_for_non_static = gpt.generate_response(vAR_joined_topics)
                                            #response_for_non_static = gpt.summarizing(response_for_non_static)
                                            response_for_conc = gpt.generate_response(vAR_conclusion)
                                            headx=vAR_new_doc.add_heading(topics)
                                            headx.style.font.color.rgb = RGBColor(0, 0, 139)
                                            headx.style.font.size = Pt(14)
                                            headx.style.font.bold = True
                                            headx.style.font.all_caps = True
                                            vAR_new_doc.add_paragraph(response_for_non_static)
                                    with col2:
                                        st.info("Formatting Content and Writing Content in a Word Document") 
                                        
                                    vAR_new_doc.add_heading("Conclusion")
                                    vAR_new_doc.add_paragraph(response_for_conc)
                                with col2:
                                    st.markdown("")
                                    bio = io.BytesIO()
                                    vAR_new_doc.save(bio)
                                    st.download_button(
                                        label="Download",
                                        data=bio.getvalue(),
                                        file_name=vAR_prompt_file+".docx",
                                        mime="docx"
                                    )
            if vAR_Model =="GPT-3.5":
                with col2:
                    if vAR_prompt_file !="":
                        st.markdown("")
                        button_placeholder = st.empty()
                        button_clicked = False
                        key=0
                        while not button_clicked:
                            key=key+1
                            button_clicked = button_placeholder.button('Generate Content',key=key)
                            break
                        if button_clicked:
                            with col2:
                                if vAR_prompt_file !="":
                                    button_placeholder.empty()
                                    # with col2:
                                    #     st.info("Writing content for the respective subtopic in the word file")
                                    vAR_new_doc = GPTTurbo.contnet_gen_turbo(vAR_prompt_file,vAR_new_doc)
                                    with col2:
                                        st.markdown("")
                                        bio = io.BytesIO()
                                        vAR_new_doc.save(bio)
                                        st.download_button(
                                            label="Download",
                                            data=bio.getvalue(),
                                            file_name=vAR_prompt_file+".docx",
                                            mime="docx"
                                        )
