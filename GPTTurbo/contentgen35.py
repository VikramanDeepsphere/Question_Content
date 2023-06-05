import docx
from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, ns
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import GPTTurbo.Topics_prompt as topic_gen
import GPTTurbo.content_prompt as content_gen
import GPTTurbo.conclusion_prompt as conclusion_gen
import GPTTurbo.Turbo35 as GPT_35
import streamlit as st

def contnet_gen_turbo(vAR_prompt_file,vAR_new_doc):
    st.info("Generating subtopics for the given topic")
    vAR_Message_log_topics = topic_gen.Sub_topics_turbo(vAR_prompt_file)
    st.info("Splitting the subtopics from the response")
    vAR_topics = GPT_35.generate_response4(vAR_Message_log_topics)
    vAR_splitted_list = vAR_topics.split("\n")
    st.info("Generating contents for the respective subtopic")
    for j in vAR_splitted_list:
        if j=="":
            vAR_splitted_list.remove(j)
    for i in vAR_splitted_list:
        if i[0].isnumeric():
            vAR_Message_log_content = content_gen.Sub_topics_turbo(i)
            vAR_content_Gen = GPT_35.generate_response4(vAR_Message_log_content)
            headx=vAR_new_doc.add_heading(i)
            headx.style.font.color.rgb = RGBColor(0, 0, 139)
            headx.style.font.size = Pt(14)
            headx.style.font.bold = True
            headx.style.font.all_caps = True
            vAR_new_doc.add_paragraph(vAR_content_Gen)
    vAR_Message_log_conclusion = conclusion_gen.conclusion_topics_turbo(vAR_prompt_file)
    vAR_Conclusion = GPT_35.generate_response4(vAR_Message_log_conclusion)
    vAR_new_doc.add_heading("Conclusion")
    vAR_new_doc.add_paragraph(vAR_Conclusion)
    return vAR_new_doc