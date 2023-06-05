import streamlit as st
from streamlit_option_menu import option_menu
import os
import io
import docx 
import source.title_1 as head
import QuestionGeneration.text_file as txt
import QuestionGeneration.split_prompt as spilt
import source.GPT_3 as gpt
import QuestionGeneration.pdf as pdf
import QuestionGeneration.pdf_textretrive as pdf_retrive
import openai
from docx import document
import QuestionGeneration.word_doc_que_gen as word_que
import GPTTurbo.final_prompt_que as prompt_que
def prev():
    st.session_state['preview']="No"
    st.session_state['preview2']="No"
def quens_gen():
    head.title()
    st.markdown("<p style='text-align: center; color: black; font-size:20px;'><span style='font-weight: bold'>Problem Statement: </span>ChatGPT Powered Business Applications for Question Generation </p>", unsafe_allow_html=True)
    st.markdown("<hr style=height:2.5px;background-color:gray>",unsafe_allow_html=True)
    w1,col1,col2,w2=st.columns((0.5,2,2.5,0.7))
    w12,col11,col22,w22=st.columns((0.5,2,2.5,0.7))
    cc2,cc1,cc3=st.columns((1,4,1))

    openai.api_key = os.environ["API_KEY"]
    with col1:
        st.markdown("### ")
        st.write('# ')
        st.markdown("<p style='text-align: left; color: black; font-size:20px;'><span style='font-weight: 600'>Select Model</span></p>", unsafe_allow_html=True)
        st.markdown("### ")
        st.markdown(" ")
        
        #st.write("# Input Typeㅤㅤㅤ")   
    with col2:
        st.markdown("## ")
        vAR_Model = st.radio("",["GPT-3","GPT-3.5",],horizontal=True,index=0)
        if vAR_Model != "Select":
            vAR_input_file_select = st.selectbox("",["Select","Text File Upload","PDF File Upload","User Enter Text","Web URL"],key="Clear_quetype")
            with col1:
                st.markdown("<p style='text-align: left; color: black; font-size:20px;'><span style='font-weight: 600'>Model Input Type</span></p>", unsafe_allow_html=True)
            # file operations
            if vAR_input_file_select == "Text File Upload" or vAR_input_file_select == "PDF File Upload":
                with col1:
                    st.write("# ")
                    st.markdown("## ")
                    st.markdown("<p style='text-align: left; color: black; font-size:20px;'><span style='font-weight: 600'>Model Input (File Upload)</span></p>", unsafe_allow_html=True)
                with col2:
                    prompt_file_txt = st.file_uploader("", type=['pdf', 'txt'])
                if 'txt' in str(prompt_file_txt):
                    vAR_content=prompt_file_txt.getvalue().decode("utf-8")
                    with col11:
                        st.markdown("# ")
                        st.markdown("<p style='text-align: left; color: black; font-size:20px;'><span style='font-weight: 600'>Preview</span></p>", unsafe_allow_html=True)
                    with col22:
                        vAR_preview = st.selectbox("",["Select","Yes","No"],key="preview")
                    if vAR_preview =="Yes":
                        with cc1:
                            st.write("")
                            st.write(vAR_content)
                    with col11:
                        st.markdown("# ")
                        st.markdown("<p style='text-align: left; color: black; font-size:20px;'><span style='font-weight: 600'>Select Question Type </span></p>", unsafe_allow_html=True)
                    with col22:
                        vAR_input_quetype = st.selectbox("",["Select","Fill in the blanks","Multiple choices","True or False","Match the following"],key="Clear_inputfile",on_change=prev)   
                    if prompt_file_txt:
                        txt_content = prompt_file_txt.getvalue().decode("utf-8")
                        to_split_content=spilt.fullstr(txt_content)
                        os.remove("Result//newfile.txt")
                        if vAR_input_quetype !="Select":
                            with col22:    
                                st.markdown("")
                                button_placeholder = st.empty()
                                button_clicked = False
                                key=0
                                while not button_clicked:
                                    key=key+1
                                    button_clicked = button_placeholder.button('Submit',key=key)
                                    break
                                if button_clicked:
                                    button_placeholder.empty()
                                    vAR_new_doc = docx.Document()
                                    fin_txt = ""
                                    with col22:
                                        st.info("Extracting Text from the Input File")
                                        st.info("Generating "+vAR_input_quetype +" Questions")
                                    no_iter=len(to_split_content)
                                    half=no_iter//2
                                    vAR_x = [(list(range(i, i+5))) for i in range(1,((no_iter*5)+1),5)]
                                    serial=0
                                    for split_content in to_split_content:
                                        if split_content==to_split_content[half]:
                                            with col22:
                                                st.info("50"+"%" +" of the Questions are Generated")
                                        serial_no = vAR_x[serial]
                                        serial+=1
                                        txt_prompt = txt.txt_file(split_content,vAR_input_quetype,serial_no)
                                        if vAR_Model =="GPT-3":
                                            vAR_responce = gpt.generate_response3(txt_prompt)
                                            fin_txt =fin_txt + str(vAR_responce)
                                        elif vAR_Model =="GPT-3.5":
                                            if vAR_input_quetype == "Fill in the blanks":
                                                vAR_responce = prompt_que.fillup(split_content,serial_no)
                                                fin_txt =fin_txt +"\n"+ str(vAR_responce)
                                            elif vAR_input_quetype == "Multiple choices":
                                                vAR_responce = prompt_que.mcq(split_content,serial_no)
                                                fin_txt =fin_txt +"\n"+ str(vAR_responce)
                                            elif vAR_input_quetype == "True or False":
                                                vAR_responce = prompt_que.trf(split_content,serial_no)
                                                fin_txt =fin_txt +"\n"+ str(vAR_responce)
                                            elif vAR_input_quetype == "Match the following":
                                                vAR_responce = gpt.generate_response3(txt_prompt)
                                                fin_txt =fin_txt + str(vAR_responce)
                                            else:
                                                pass
                                        else:
                                            pass
                                    with col22:
                                        st.info("Creates a Word Document with Questions ")
                                    vAR_final_docs = word_que.heading_content_que_gen(vAR_new_doc,fin_txt,vAR_input_quetype)
                                    #vAR_final_docs.save(str(vAR_input_quetype)+".docx")
                                    bio = io.BytesIO()
                                    vAR_final_docs.save(bio)
                                    # with bc2:
                                    #     st.markdown("")
                                    #     st.button("Clear", on_click=cr.clear_text)
                                    with col22:
                                        st.markdown("")
                                        st.download_button(
                                            label="Download",
                                            data=bio.getvalue(),
                                            file_name=vAR_input_quetype+".docx",
                                            mime="docx"
                                        )
                                    
    ################################################################################################################            
                # pdf file operatons 
                if 'pdf' in str(prompt_file_txt):
                    pdf.header_footer_cuter(prompt_file_txt)
                    pdf_content = pdf_retrive.pdf_text()
                    with col11:
                        st.markdown("# ")
                        st.markdown("<p style='text-align: left; color: black; font-size:20px;'><span style='font-weight: 600'>Preview</span></p>", unsafe_allow_html=True)
                    with col22:
                        vAR_preview = st.selectbox("",["Select","Yes","No"],key="preview2")
                    if vAR_preview =="Yes":
                        with cc1:
                            st.write("")
                            # pdf.header_footer_cuter(prompt_file_txt)
                            # pdf_content = pdf_retrive.pdf_text()
                            st.write(pdf_content)
                    with col11:
                        st.markdown("# ")
                        st.markdown("<p style='text-align: left; color: black; font-size:20px;'><span style='font-weight: 600'>Select Question Type </span></p>", unsafe_allow_html=True)
                    with col22:
                        vAR_input_quetype = st.selectbox("",["Select","Fill in the blanks","Multiple choices","True or False","Match the following"],key="Clear_inputfile",on_change=prev) 
                    if vAR_input_quetype !="Select":
                        with col22:
                            st.markdown("")
                            button_placeholder = st.empty()
                            button_clicked = False
                            key=0
                            while not button_clicked:
                                key=key+1
                                button_clicked = button_placeholder.button('Submit',key=key)
                                break
                            if button_clicked:
                                button_placeholder.empty()
                                with col22:
                                    st.info("Removing header and footer from the Input File")
                                    
                                if prompt_file_txt is not None:
                                    with col22:
                                        st.info("Extracting Text from the Input File")
                                    to_split_content=spilt.fullstr(pdf_content)
                                    os.remove("Result//newfile.txt")
                                    with col22:
                                        st.info("Summarizing the content and passes to the model")
                                    #Creating a new word document.
                                    vAR_new_doc = docx.Document()
                                    fin_txt = ""
                                    no_iter=len(to_split_content)
                                    half=no_iter//2
                                    with col22:
                                        st.info("Generating "+vAR_input_quetype+" Questions")
                                    vAR_x = [(list(range(i, i+5))) for i in range(1,((no_iter*5)+1),5)]
                                    serial=0
                                    for split_content in to_split_content:
                                        if split_content==to_split_content[half]:
                                            with col22:
                                                st.info("50"+"%"+" of the Questions are Generated")
                                        split_content = gpt.summarizing(split_content)
                                        serial_no = vAR_x[serial]
                                        serial+=1
                                        txt_prompt = txt.txt_file(split_content,vAR_input_quetype,serial_no)
                                        if vAR_Model =="GPT-3":
                                            vAR_responce = gpt.generate_response3(txt_prompt)
                                            fin_txt =fin_txt + str(vAR_responce)
                                        elif vAR_Model =="GPT-3.5":
                                            if vAR_input_quetype == "Fill in the blanks":
                                                vAR_responce = prompt_que.fillup(split_content,serial_no)
                                                fin_txt =fin_txt +"\n"+ str(vAR_responce)
                                            elif vAR_input_quetype == "Multiple choices":
                                                vAR_responce = prompt_que.mcq(split_content,serial_no)
                                                fin_txt =fin_txt +"\n"+ str(vAR_responce)
                                            elif vAR_input_quetype == "True or False":
                                                vAR_responce = prompt_que.trf(split_content,serial_no)
                                                fin_txt =fin_txt +"\n"+ str(vAR_responce)
                                            elif vAR_input_quetype == "Match the following":
                                                vAR_responce = gpt.generate_response3(txt_prompt)
                                                fin_txt =fin_txt + str(vAR_responce)
                                            else:
                                                pass
                                        else:
                                            pass
                                    with col22:
                                        st.info("Creates a Word Document with Questions")
                                    vAR_final_docs = word_que.heading_content_que_gen(vAR_new_doc,fin_txt,vAR_input_quetype)
                                    with col22:
                                        st.markdown("")
                                        bio = io.BytesIO()
                                        vAR_final_docs.save(bio)
                                        st.download_button(
                                            label="Download",
                                            data=bio.getvalue(),
                                            file_name=vAR_input_quetype+".docx",
                                            mime="docx"
                                        )
    ##################################################################################################################                              
            # Passage operations
            elif vAR_input_file_select == "User Enter Text" or vAR_input_file_select == "Web URL":
                with col1:
                    st.markdown("# ")
                    st.markdown("# ")
                    st.markdown("<p style='text-align: left; color: black; font-size:20px;'><span style='font-weight: 600'>Model Input (Enter Text)</span></p>", unsafe_allow_html=True)
                with col1:
                    st.markdown("## ")
                    st.markdown("# ")
                    st.markdown("<p style='text-align: left; color: black; font-size:20px;'><span style='font-weight: 600'>Select Question Type</span></p>", unsafe_allow_html=True)
                vAR_passage = st.text_area("",key = "Clear_input_para")
                to_split_content=spilt.fullstr(vAR_passage)
                with col2:
                        vAR_input_quetype = st.selectbox("",["Select","Fill in the blanks","Multiple choices","True or False","Match the following"],key="Clear_inputfile") 
                if vAR_input_quetype !="Select":
                    with col22:    
                        st.markdown("# ")
                        button_placeholder = st.empty()
                        button_clicked = False
                        key=0
                        while not button_clicked:
                            key=key+1
                            button_clicked = button_placeholder.button('Submit',key=key)
                            break
                        if button_clicked:
                            if vAR_passage !="":
                                if len(vAR_passage) >= 500:
                                    button_placeholder.empty()
                                    #Creating a new word document.
                                    vAR_new_doc = docx.Document()
                                    fin_txt = ""
                                    no_iter=len(to_split_content)
                                    half=no_iter//2
                                    with col22:
                                        st.info("Generating "+vAR_input_quetype+" Questions")
                                    vAR_x = [(list(range(i, i+5))) for i in range(1,((no_iter*5)+1),5)]
                                    serial=0
                                    for split_content in to_split_content:
                                        if split_content==to_split_content[half]:
                                            with col22:
                                                st.info("50"+"%"+" of the Questions are Generated")
                                        serial_no = vAR_x[serial]
                                        serial += 1
                                        txt_prompt = txt.txt_file(split_content,vAR_input_quetype,serial_no)
                                        if vAR_Model =="GPT-3":
                                            vAR_responce = gpt.generate_response3(txt_prompt)
                                            fin_txt =fin_txt + str(vAR_responce)
                                        elif vAR_Model =="GPT-3.5":
                                            if vAR_input_quetype == "Fill in the blanks":
                                                vAR_responce = prompt_que.fillup(split_content,serial_no)
                                                fin_txt =fin_txt +"\n"+ str(vAR_responce)
                                            elif vAR_input_quetype == "Multiple choices":
                                                vAR_responce = prompt_que.mcq(split_content,serial_no)
                                                fin_txt =fin_txt +"\n"+ str(vAR_responce)
                                            elif vAR_input_quetype == "True or False":
                                                vAR_responce = prompt_que.trf(split_content,serial_no)
                                                fin_txt =fin_txt +"\n"+ str(vAR_responce)
                                            elif vAR_input_quetype == "Match the following":
                                                vAR_responce = gpt.generate_response3(txt_prompt)
                                                fin_txt =fin_txt + str(vAR_responce)
                                            else:
                                                pass
                                        else:
                                            pass
                                    vAR_final_docs = word_que.heading_content_que_gen(vAR_new_doc,fin_txt,vAR_input_quetype)
                                    with col22:
                                        #st.write("# ")
                                        st.info("Creates a Word Document with Questions")
                                        bio = io.BytesIO()
                                        vAR_final_docs.save(bio)
                                    with col22:
                                        st.markdown("")
                                        st.download_button(
                                            label="Download",
                                            data=bio.getvalue(),
                                            file_name=vAR_input_quetype+".docx",
                                            mime="docx"
                                        )
                                else:
                                    with col2:
                                        st.warning("The Content is too Short")
                            else:
                                with col2:
                                    st.warning("Enter the Passage")
